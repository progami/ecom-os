from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

src = Path('Bussiness Knowledge.docx')
if not src.exists():
    raise SystemExit('Source DOCX not found: %s' % src)

# Backup once if not already present
backup = Path('Bussiness Knowledge.backup.before-rewrite.docx')
if not backup.exists():
    backup.write_bytes(src.read_bytes())

D = Document()

# Base style
style = D.styles['Normal']
style.font.name = 'Garamond'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Garamond')
style.font.size = Pt(12)

def addp(text, bold=False):
    p = D.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

# Header block
addp('September 12, 2025')
addp('USCIS')
addp('Attn: Premium I-129 E-1/E-2/E-2C (Box 88781)')
addp('131 S. Dearborn St., 3rd Floor')
addp('Chicago, IL 60603-5517')
D.add_paragraph('')

addp('Petitioner:          Targon LLC')
addp('Beneficiary:         Jarrar Amjad')
addp('Nationality:         Pakistan')
addp('Ownership & Control: 100% owner; will develop and direct as CEO')
addp('Re: I-129 Petition — E-2 Treaty Investor (Change of Status)')
D.add_paragraph('')

addp('Respected Director,')

addp('Please accept this letter in support of my I-129 petition to classify me as an E-2 Treaty Investor. I own 100% of Targon LLC and will develop and direct the enterprise as CEO. The company manufactures and distributes drop cloths and protective sheeting under the Caelum Star brand, launching on Amazon in October 2025 and expanding to Walmart thereafter.')

# Source and path of funds
addp('Source and Path of Investment Funds (Gift)', bold=True)
addp('My investment funds were lawfully obtained through a bona fide financial gift from Ammar Amjad, an employee of Apple Inc. His salary and relocation allowances are deposited into his Bank of America account ending in 8857 and were transferred to Targon LLC’s business accounts at JPMorgan Chase (ending in 1203 and 9899). A notarized gift affidavit and corresponding bank statements tracing the path of funds are enclosed (App. X-1, X-2).')

# Transfer details list
transfer = [
    '04/25/2025 — $10,000.00 (with $8,898.25 returned; net gift $1,101.75)',
    '08/26/2025 — $8,000.00 (full gift)',
    '08/29/2025 — $20,000.00 (full gift)',
    '09/02/2025 — $20,700.00 (full gift)',
    '09/05/2025 — $31,300.00 (full gift)',
]
D.add_paragraph('Transfer Details:')
for item in transfer:
    p = D.add_paragraph(style='List Bullet')
    p.add_run(item)
addp('Total Gifted: $81,101.75')
addp('This was a voluntary, unconditional gift. The donor retains no claim, lien, or interest in the funds.')

# Investment at risk
addp('Investment Is Substantial and At Risk', bold=True)
addp('Total investment committed (at risk): $81,102; $53,414.58 (65.9%) deployed as of September 10, 2025. The remaining balance is reserved for U.S. import tariffs at port and inbound freight charges on the arriving shipment.')

# Real and operating
addp('Real and Operating Commercial Enterprise', bold=True)
addp('Inventory is manufactured and arriving in the United States on September 29, 2025. Amazon Seller Central is active and verified; Walmart marketplace is approved. Property insurance is secured; the commercial lease is pending signature on September 15, 2025; and a 3PL warehousing agreement is executed (Apps C-4, C-8–C-9, D-10, E-1–E-6). Brand and IP evidence (including Caelum Star and USPTO records) are included in App. C-7.')

# Jobs / non-marginality
addp('Non-Marginality and U.S. Job Creation', bold=True)
addp('One full-time U.S. Operations Manager was hired in August 2025; one part-time Associate will be hired within 30 days of inventory arrival. The company commits to a minimum of nine U.S. jobs by 2030 (six full-time, three part-time), excluding the founder. Financial projections support sustained payroll and operating expenses from Year 2 onward (App. B).')

# Role and qualifications
addp('Applicant’s Role and Qualifications', bold=True)
addp('I will develop and direct Targon LLC as CEO, leveraging four years of prior success as COO in the United Kingdom in the same product category, as well as an AWS Certified Solutions Architect – Associate credential. Detailed duties are outlined below and in the business plan.')

# Duties
addp('Executive Duties', bold=True)
for duty in [
    'Conceptualize and lead expansion initiatives; direct feasibility, operations, and financial management.',
    'Direct and coordinate Targon LLC activities to optimize efficiency and profitability.',
    'Lead market development and promotion to increase share and competitive position.',
    'Review operations and financial performance; forecast production and costs against objectives.',
    'Appoint managers and delegate responsibilities; oversee performance and compliance.',
    'Negotiate and approve supplier and vendor contracts.',
    'Engage clients and local officials to drive profitability and ensure regulatory compliance.',
]:
    p = D.add_paragraph(style='List Bullet')
    p.add_run(duty)

# Intent to depart
addp('Intent to Depart', bold=True)
addp('I intend to remain in the United States while in valid nonimmigrant status and will depart upon its expiration.')

# Supporting documents
addp('Supporting Documents', bold=True)
for line in [
    'Certificate of Formation and EIN of Targon LLC',
    'Lease agreement with proof of rent payment and corresponding bank statement',
    'Operating Agreement — Targon LLC',
    'Licenses & Permits of Targon LLC',
    'Liability insurance of Targon LLC',
    'Form I-9, SSN card, state ID, and paystubs with corresponding bank statements (Targon LLC)',
    '3PL contract providing warehouse space (SyzTech Logistics)',
    'Two supplier contracts with invoices and corresponding bank statements',
    'Purchase and sale invoices with corresponding business bank statements (ending 1203)',
    'Amazon & Walmart account dashboards',
    'Investment breakdown',
    'Comprehensive business plan with five-year projections',
    'Photographs of Targon LLC',
    'Appendices index: App. A–G (business history, financials, legal/corporate, supplier/3PL, platforms, product, organization) and App. X-1–X-2 (gift and bank statements/invoices)'
]:
    p = D.add_paragraph(style='List Bullet')
    p.add_run(line)

# Closing
D.add_paragraph('')
addp('Thank you for your consideration of my E-2 change of status application.')
D.add_paragraph('')
addp('Sincerely,')
addp('Jarrar Amjad')
addp('Founder & CEO, Targon LLC')

D.save(str(src))
print('Updated DOCX:', src)

